import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Definir márgenes de 1 pulgada
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Estilo de Portada
title = doc.add_paragraph()
title_run = title.add_run("MANUAL DEL USUARIO")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(212, 175, 55) # Oro #D4AF37
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
sub_run = subtitle.add_run("Sistema de Gestión de Restaurante: \"El Pulpazo\" (v3.1)")
sub_run.font.size = Pt(16)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(100, 100, 100)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

def add_custom_heading(text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(212, 175, 55)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(40, 40, 40)
        run.font.bold = True
    return h

def add_image_safe(image_filename, caption=""):
    img_path = os.path.abspath(os.path.join("evidencias_screenshots", image_filename))
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.8))
        if caption:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cap = p_cap.add_run(f"Figura: {caption}")
            r_cap.font.size = Pt(9)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(120, 120, 120)

# 1. Portada y Control de Versiones
add_custom_heading("1. Portada y Control de Versiones", 1)
table = doc.add_table(rows=6, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
pv_data = [
    ("Nombre del Sistema:", "El Pulpazo — Sistema de Gestión de Restaurante y Marisquería"),
    ("Versión del Manual:", "v3.1 (Versión Mejorada con Cocina, Recetas, Descuentos y Reportes)"),
    ("Fecha de Actualización:", "23 de Julio de 2026"),
    ("Asignatura:", "Tópicos de calidad para el desarrollo de software (Actividad 5)"),
    ("Equipo de Desarrollo:", "Lopez Estrella B. O., Castro Nuñez N. M., Pech Ake S. A., May de los Santos J. J."),
    ("Público Objetivo:", "Administradores de restaurante, cajeros, meseros, personal de cocina y operaciones.")
]
for row_idx, (k, v) in enumerate(pv_data):
    cell_k = table.cell(row_idx, 0)
    cell_k.text = k
    cell_k.paragraphs[0].runs[0].font.bold = True
    shading = parse_xml(r'<w:shd {} w:fill="F4F4F4"/>'.format(nsdecls('w')))
    cell_k._tc.get_or_add_tcPr().append(shading)
    table.cell(row_idx, 1).text = v

doc.add_paragraph()

# 2. Introducción y Objetivo
add_custom_heading("2. Introducción y Objetivo", 1)
doc.add_paragraph(
    "El Pulpazo v3.1 es una aplicación web integral diseñada para optimizar la operación diaria de restaurantes "
    "y marisquerías. Permite administrar en tiempo real el plano de mesas, la toma de comandas, la preparación en "
    "cocina estilo Kanban, el cobro con desglose del IVA del 16%, descuentos por cupón, división de cuentas, "
    "inventario automático por recetas y analítica ejecutiva. El objetivo de este manual es guiar de forma "
    "clara, sencilla y coloquial a cualquier miembro del personal en el uso correcto de la plataforma."
)

p_link = doc.add_paragraph()
r_link = p_link.add_run("Acceso Web Inmediato: https://santiagopech490-alt.github.io/pescaderia_V3/")
r_link.font.bold = True
r_link.font.color.rgb = RGBColor(0, 102, 204)

# 3. Índice
add_custom_heading("3. Índice / Mapa de Contenidos", 1)
indice_items = [
    "1. Portada y Control de Versiones", "2. Introducción y Objetivo", "3. Índice / Mapa de Contenidos",
    "4. Requisitos del Sistema", "5. Convenciones e Iconografía", "6. Inicio de Sesión",
    "7. Descripción General del Sistema", "8. Roles y Perfiles de Usuario", "9. Módulos y Operaciones del Usuario",
    "10. Validaciones Importantes", "11. Atajos y Accesos Rápidos", "12. Solución de Problemas (Troubleshooting)",
    "13. Preguntas Frecuentes (FAQ)", "14. Buenas Prácticas", "15. Seguridad y Confidencialidad de Datos",
    "16. Canal de Soporte y Contacto", "17. Glosario y Conclusiones"
]
for item in indice_items:
    p_item = doc.add_paragraph(style='List Bullet')
    p_item.add_run(item)

# 4. Requisitos del Sistema
add_custom_heading("4. Requisitos del Sistema", 1)
req_table = doc.add_table(rows=5, cols=2)
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
req_table.style = 'Table Grid'
req_headers = ["Componente", "Requisito Mínimo Recomendado"]
for i, h_t in enumerate(req_headers):
    cell = req_table.cell(0, i)
    cell.text = h_t
    cell.paragraphs[0].runs[0].font.bold = True
    shd = parse_xml(r'<w:shd {} w:fill="D4AF37"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shd)

req_data = [
    ("Navegador Web", "Google Chrome v100+, Microsoft Edge v100+, Safari v15+ o Mozilla Firefox v100+."),
    ("Sistema Operativo", "Windows 10/11, macOS, Linux, Android (Tablet/Celular) o iOS/iPadOS."),
    ("Conexión a Internet", "Estable (Mínimo 2 Mbps) para sincronización en la nube con Supabase."),
    ("Permisos de Cuenta", "Credenciales de usuario asignadas por el Administrador según su rol.")
]
for r_idx, (c1, c2) in enumerate(req_data, start=1):
    req_table.cell(r_idx, 0).text = c1
    req_table.cell(r_idx, 1).text = c2

doc.add_paragraph()

# 5. Convenciones
add_custom_heading("5. Convenciones e Iconografía", 1)
doc.add_paragraph(
    "• ⚠️ Advertencia: Información crítica para prevenir errores operativos o pérdida de datos.\n"
    "• 💡 Tip / Sugerencia: Consejos útiles para trabajar de forma más rápida y eficiente.\n"
    "• 🔒 Requiere Permisos: Acción restringida únicamente para usuarios con roles autorizados.\n"
    "• 🟢 Mesa Libre: Disponible para asignar a clientes en el comedor.\n"
    "• 🔴 Mesa Ocupada: Comanda activa con tiempo transcurrido.\n"
    "• 🟠 Mesa Reservada: Apartada para comensales futuros.\n"
    "• 🔔 Notificación Flotante: Alerta de stock bajo o comanda lista en cocina."
)

# 6. Inicio de Sesión
add_custom_heading("6. Inicio de Sesión", 1)
doc.add_paragraph(
    "1. Accede a la URL oficial del sistema: https://santiagopech490-alt.github.io/pescaderia_V3/\n"
    "2. Ingresa tu correo electrónico registrado y tu contraseña.\n"
    "3. Haz clic en el botón amarillo 'INICIAR SESIÓN'."
)
add_image_safe("01_login_autenticacion.png", "Pantalla de Inicio de Sesión de El Pulpazo")

# 7. Descripción General
add_custom_heading("7. Descripción General del Sistema", 1)
doc.add_paragraph(
    "El sistema cuenta con una interfaz limpia en modo oscuro con acentos dorados (#D4AF37). Su estructura consta de:\n"
    "1. Barra Lateral (Sidebar): Menú para cambiar entre módulos autorizados.\n"
    "2. Encabezado (Header): Muestra iniciales, rol (Administrador, Cajera, Mesero, Cocina), campanita de notificaciones, selector Sol/Luna y Cerrar Sesión.\n"
    "3. Área Central: Visualización interactiva del módulo activo."
)

# 8. Roles
add_custom_heading("8. Roles y Perfiles de Usuario", 1)
roles_table = doc.add_table(rows=5, cols=5)
roles_table.alignment = WD_TABLE_ALIGNMENT.CENTER
roles_table.style = 'Table Grid'
roles_headers = ["Rol", "Correo de Acceso", "Contraseña", "Vistas Visibles", "Permisos y Capacidades"]
for i, h_t in enumerate(roles_headers):
    cell = roles_table.cell(0, i)
    cell.text = h_t
    cell.paragraphs[0].runs[0].font.bold = True
    shd = parse_xml(r'<w:shd {} w:fill="D4AF37"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shd)

roles_data = [
    ("Administrador", "admin@elpulpazo.com", "admin@elpulpazo", "Todas las pantallas", "Control total: mapa de mesas, platillos, recetas, inventario, reportes CSV y respaldos."),
    ("Cajera / Cajero", "Cajero@Pulpazo.com", "Cajero", "Mesas, Pedidos, Carrito, Inventario, Corte, Descuentos", "Procesar cobros, calcular cambio/IVA, emitir tickets digitales, registrar gastos/abonos."),
    ("Mesero", "mesero@elpulpazo.com", "mesero123", "Mesas, Menú, Carrito, Historial", "Consultar disponibilidad de mesas, tomar pedidos en el menú digital y enviar al carrito."),
    ("Cocina", "cocina@elpulpazo.com", "cocina123", "Vista Cocina, Historial", "Pantalla Kanban de preparación de platillos (Pendiente, En preparación, Listo).")
]
for r_idx, r_d in enumerate(roles_data, start=1):
    for c_idx, val in enumerate(r_d):
        roles_table.cell(r_idx, c_idx).text = val

doc.add_paragraph()

# 9. Módulos
add_custom_heading("9. Módulos y Operaciones del Usuario", 1)

add_custom_heading("9.1 Plano Interactivo de Mesas", 2)
doc.add_paragraph("Permite visualizar la distribución física del salón de comedor en tiempo real.")
add_image_safe("02_plano_de_mesas.png", "Plano Interactivo de Mesas")

add_custom_heading("9.2 Menú Digital y Tomar Pedido", 2)
doc.add_paragraph("Navegación por categorías de alimentos (Entradas, Platos Fuertes, Bebidas, Postres).")
add_image_safe("03_menu_platillos.png", "Menú Digital de Platillos")

add_custom_heading("9.3 Carrito, Cobros, IVA y División de Cuenta", 2)
doc.add_paragraph("Selección de mesa, método de pago, cupones, desglose del IVA 16% e importe total.")
add_image_safe("04_carrito_checkout.png", "Carrito de Compras y Cobro")

add_custom_heading("9.4 Vista de Cocina Kanban (NUEVO v3.1)", 2)
doc.add_paragraph("Organizada en 3 columnas: Pendiente, En preparación y Listo.")
add_image_safe("09_vista_cocina_kanban.png", "Vista de Cocina Kanban")

add_custom_heading("9.5 Control de Inventario Automático y Recetas", 2)
doc.add_paragraph("Descuento automático de ingredientes por receta al confirmar ventas.")
add_image_safe("05_inventario_stock.png", "Control de Inventario")

add_custom_heading("9.6 Cupones y Descuentos (NUEVO v3.1)", 2)
doc.add_paragraph("Gestión de cupones de descuento (% o monto fijo).")
add_image_safe("11_cupones_descuentos.png", "Cupones y Descuentos")

add_custom_heading("9.7 Reportes de Ventas y Exportación CSV (NUEVO v3.1)", 2)
doc.add_paragraph("Analítica avanzada con exportación a Excel / CSV.")
add_image_safe("10_reportes_ventas.png", "Reportes de Ventas")

add_custom_heading("9.8 Dashboard Ejecutivo", 2)
add_image_safe("06_dashboard_ejecutivo.png", "Dashboard Ejecutivo")

add_custom_heading("9.9 Administración de Platillos", 2)
add_image_safe("07_panel_administracion.png", "Panel Admin Platillos")

# 10. Validaciones Importantes
add_custom_heading("10. Validaciones Importantes", 1)
val_table = doc.add_table(rows=7, cols=3)
val_table.alignment = WD_TABLE_ALIGNMENT.CENTER
val_table.style = 'Table Grid'
val_headers = ["Campo / Módulo", "Regla de Negocio", "Ejemplo / Resultado"]
for i, h_t in enumerate(val_headers):
    cell = val_table.cell(0, i)
    cell.text = h_t
    cell.paragraphs[0].runs[0].font.bold = True
    shd = parse_xml(r'<w:shd {} w:fill="F4F4F4"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shd)

val_data = [
    ("Mesa en Checkout", "Obligatorio seleccionar una mesa antes de cobrar.", "Si no se elige mesa, no se procesa el pago."),
    ("Cantidad en Carrito", "No se permiten cantidades negativas o cero.", "El botón '-' elimina el ítem al llegar a 0."),
    ("Cálculo de IVA", "Aplica tasa fija del 16% sobre el subtotal neto.", "Subtotal = $100.00 → IVA = $16.00 → Total = $116.00."),
    ("Cupones", "Se descuenta antes del cálculo de IVA.", "Subtotal $200 - Cupón $20 = $180 netos + $28.80 IVA."),
    ("Nivel de Stock", "Alerta roja automática al caer el stock.", "Stock Actual (1.2 kg) < Stock Mínimo (2.0 kg) → 'Bajo Stock'."),
    ("Acceso a Rutas", "Restringido estrictamente por el rol asignado.", "El rol Cocina solo accede a /cocina e /historial.")
]
for r_idx, r_d in enumerate(val_data, start=1):
    for c_idx, val in enumerate(r_d):
        val_table.cell(r_idx, c_idx).text = val

doc.add_paragraph()

# 11. Atajos
add_custom_heading("11. Atajos y Accesos Rápidos", 1)
doc.add_paragraph(
    "• Campanita Notificaciones: Hacer clic en la campanita flotante del header.\n"
    "• Cambiar Tema Visual: Clic en el icono Sol/Luna en el encabezado.\n"
    "• Colapsar / Abrir Menú: Clic en el icono de Hamburguesa (≡) junto al logo.\n"
    "• Búsqueda de Platillos: Usar la barra de búsqueda en 'Tomar Pedido'.\n"
    "• Cerrar Sesión: Clic en el icono de salir en la esquina superior derecha."
)

# 12. Solución de Problemas
add_custom_heading("12. Solución de Problemas (Troubleshooting)", 1)
tb_table = doc.add_table(rows=5, cols=3)
tb_table.alignment = WD_TABLE_ALIGNMENT.CENTER
tb_table.style = 'Table Grid'
tb_headers = ["Síntoma", "Causa Probable", "Solución Recomendada"]
for i, h_t in enumerate(tb_headers):
    cell = tb_table.cell(0, i)
    cell.text = h_t
    cell.paragraphs[0].runs[0].font.bold = True
    shd = parse_xml(r'<w:shd {} w:fill="F4F4F4"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shd)

tb_data = [
    ("No puedo iniciar sesión.", "Credenciales mal escritas o usuario no registrado.", "Verifica mayúsculas o solicita revisión al Administrador."),
    ("Las mesas no guardaron cambios.", "No se presionó 'Salir de Edición'.", "Presiona siempre 'Salir de Edición' para guardar en Supabase."),
    ("Los platillos no cargan.", "Pérdida temporal de conexión a internet.", "Revisa tu red y presiona F5 para recargar."),
    ("Me aparece 'Acceso Restringido'.", "Intentaste entrar a una vista no permitida.", "Inicia sesión con la cuenta del rol adecuado (ej. Admin).")
]
for r_idx, r_d in enumerate(tb_data, start=1):
    for c_idx, val in enumerate(r_d):
        tb_table.cell(r_idx, c_idx).text = val

doc.add_paragraph()

# 13. FAQ
add_custom_heading("13. Preguntas Frecuentes (FAQ)", 1)
doc.add_paragraph(
    "¿Cómo divido una cuenta entre varias personas?\n"
    "En la pantalla de pago exitoso del Carrito, presiona 'Dividir Cuenta' y elige entre partes iguales, porcentaje o monto manual.\n\n"
    "¿Los platillos descuentan ingredientes automáticamente?\n"
    "Sí, al confirmar un pedido, el sistema descuenta las cantidades configuradas en la receta.\n\n"
    "¿Puedo exportar reportes a Excel?\n"
    "Sí. En la sección de Reportes, haz clic en 'Exportar CSV'."
)

# 14. Buenas Prácticas
add_custom_heading("14. Buenas Prácticas", 1)
doc.add_paragraph(
    "1. 🔒 Cierra sesión al retirarte de tu turno.\n"
    "2. 💡 Revisa la campanita de notificaciones para atender alertas de stock.\n"
    "3. 👨‍🍳 Mantén actualizada la pantalla de Cocina pasando pedidos a 'Listo'.\n"
    "4. ⚠️ Verifica los cupones antes de aplicarlos en caja."
)

# 15. Seguridad
add_custom_heading("15. Seguridad y Confidencialidad de Datos", 1)
doc.add_paragraph(
    "La plataforma utiliza Supabase (PostgreSQL en la nube). Las contraseñas están cifradas con bcrypt "
    "y las tablas implementan políticas Row Level Security (RLS) por rol para proteger la información sensible."
)

# 16. Canal de Soporte
add_custom_heading("16. Canal de Soporte y Contacto", 1)
doc.add_paragraph(
    "• Equipo de Soporte: Lopez Estrella B. O., Castro Nuñez N. M., Pech Ake S. A., May de los Santos J. J.\n"
    "• Correo de Contacto: soporte@elpulpazo.com\n"
    "• Horario de Atención: Lunes a Domingo de 8:00 AM a 10:00 PM."
)

# 17. Glosario y Conclusiones
add_custom_heading("17. Glosario y Conclusiones", 1)
doc.add_paragraph(
    "Glosario:\n"
    "• RBAC: Control de acceso basado en roles de usuario.\n"
    "• Kanban: Sistema visual de columnas (Pendiente, En preparación, Listo) para gestión de comandas.\n"
    "• Supabase: Backend as a Service con PostgreSQL en la nube.\n"
    "• CSV: Archivo de texto separado por comas compatible con Microsoft Excel.\n\n"
    "Conclusión:\n"
    "El manual del usuario de El Pulpazo v3.1 garantiza que todo el personal operativo maneje la plataforma con fluidez, "
    "llevando un control impecable del salón, la cocina, las finanzas y el inventario."
)

for filename in ["MANUAL_DEL_USUARIO.docx", "MANUAL_DEL_USUARIO_RUBRICA.docx"]:
    try:
        doc.save(filename)
        print(f"¡Archivo {filename} generado exitosamente!")
    except Exception as e:
        print(f"No se pudo guardar {filename}: {e}")
