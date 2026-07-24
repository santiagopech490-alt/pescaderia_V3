import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Definir márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Estilo de Título Principal
title = doc.add_paragraph()
title_run = title.add_run("Manual del Usuario del Producto")
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(212, 175, 55) # Oro #D4AF37
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
sub_run = subtitle.add_run("Sistema de Gestión de Restaurante: \"El Pulpazo\"")
sub_run.font.size = Pt(16)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(100, 100, 100)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph() # Espacio

# Función para agregar títulos
def add_custom_heading(text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(212, 175, 55)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(40, 40, 40)
        run.font.bold = True
    return h

# 1. Introducción
add_custom_heading("1. Introducción y Bienvenida", 1)
p = doc.add_paragraph(
    "Bienvenido al Manual del Usuario de El Pulpazo. Este sistema es una plataforma web moderna "
    "para la administración integral de marisquerías y restaurantes. Permite controlar en tiempo real "
    "la distribución de comedores, la toma de comandas, el cobro con desglose del IVA del 16%, "
    "las existencias en inventario y el análisis ejecutivo de ventas."
)
p.style.font.size = Pt(11)

p_link = doc.add_paragraph()
r_link = p_link.add_run("Acceso Web Inmediato: https://santiagopech490-alt.github.io/pescaderia_V3/")
r_link.font.bold = True
r_link.font.color.rgb = RGBColor(0, 102, 204)

# 2. Acceso y Roles
add_custom_heading("2. Acceso al Sistema y Roles Operativos", 1)
doc.add_paragraph(
    "El sistema cuenta con inicio de sesión por correo y contraseña registrados en la nube (Supabase Auth). "
    "Según el puesto del empleado, el sistema adapta sus permisos y opciones del menú:"
)

# Tabla de Roles
table = doc.add_table(rows=4, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

headers = ["Rol", "Vistas Disponibles", "Descripción de Permisos"]
for i, h_text in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h_text
    cell.paragraphs[0].runs[0].font.bold = True
    shading = parse_xml(r'<w:shd {} w:fill="D4AF37"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading)

data = [
    ("Administrador", "Todas las pantallas", "Acceso completo a mesas, edición de mapa, catálogo de platillos, inventario, reportes y finanzas."),
    ("Cajera", "Mesas, Pedidos, Carrito, Inventario", "Gestión de cobros, registro de gastos/abonos, seguimiento de comedores y visualización del stock."),
    ("Mesero", "Mesas, Menú, Carrito", "Consulta rápida de ocupación de mesas, toma de comanda y envío a cocina.")
]

for row_idx, row_data in enumerate(data, start=1):
    for col_idx, cell_value in enumerate(row_data):
        table.cell(row_idx, col_idx).text = cell_value

doc.add_paragraph()

# 3. Módulos
add_custom_heading("3. Guía Paso a Paso de los Módulos", 1)

add_custom_heading("3.1 Plano Interactivo de Mesas", 2)
doc.add_paragraph(
    "Permite visualizar la distribución física del salón de comedor en tiempo real mediante colores:\n"
    "• Verde (Libre): Mesa disponible para asignar nuevos comensales.\n"
    "• Rojo (Ocupado): Mesa con clientes consumiendo y tiempo transcurrido.\n"
    "• Naranja (Reservado): Mesa apartada para clientes futuros.\n\n"
    "Modo Edición (Administrador): Permite arrastrar mesas con el mouse/touch en la cuadrícula de 900x450, "
    "cambiar la forma (rectangular/circular), modificar dimensiones o rotación. Las posiciones se guardan de forma permanente en la nube."
)

add_custom_heading("3.2 Menú Digital y Tomar Pedido", 2)
doc.add_paragraph(
    "Organizado por categorías (Entradas, Platos Fuertes, Bebidas y Postres). Cuenta con un buscador en tiempo real "
    "para filtrar platillos por nombre y agregar ítems al pedido con el botón '+ AGREGAR'."
)

add_custom_heading("3.3 Carrito, Cobros e IVA (Checkout)", 2)
doc.add_paragraph(
    "Módulo de liquidación de cuentas. Es obligatorio seleccionar la mesa activa (ej. Mesa A1) y el método de pago (Efectivo o Tarjeta). "
    "El sistema desglosa automáticamente el Subtotal, calcula el IVA del 16% (Subtotal * 0.16) y genera un Ticket Digital de compra con folio único."
)

add_custom_heading("3.4 Control de Inventario e Insumos", 2)
doc.add_paragraph(
    "Control de ingredientes clave (Pulpo, Camarón, Pescado, Limón, Aceite) con sus unidades de medida (kg, L, g, unidades). "
    "Incluye alertas en color rojo ('Bajo Stock') cuando la cantidad cae por debajo del stock mínimo configurado."
)

add_custom_heading("3.5 Dashboard Ejecutivo y Métricas", 2)
doc.add_paragraph(
    "Muestra los indicadores clave (KPIs): Ventas diarias ($), Comandas activas, Margen operativo (%) y Ocupación (%). "
    "Incluye gráficas interactivas por hora y ranking de los platillos más vendidos."
)

try:
    doc.save("MANUAL_DEL_USUARIO.docx")
    print("¡Archivo MANUAL_DEL_USUARIO.docx guardado exitosamente!")
except PermissionError:
    doc.save("MANUAL_DEL_USUARIO_V2.docx")
    print("¡Archivo MANUAL_DEL_USUARIO_V2.docx guardado exitosamente (el original estaba abierto en Word)!")
